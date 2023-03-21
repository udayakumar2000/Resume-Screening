from flask import Flask, render_template, request, jsonify
import os
import docx2txt
import docx
import nltk
import re
import pickle
import pyresparser
import pandas as pd
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

def clean_resume(text):
    text = str(text)
    stop_words = set(stopwords.words('english'))
    text = re.sub('http\S+\s*', ' ', text) 
    text = re.sub('[%s]' % re.escape("""!"#$%&'()*+,-./:;<=>?@[\]^_`{|}~"""), ' ', text)  
    text = re.sub(r'\s+', ' ', text)

    #replacing words with numbers
    #Eg: I have one year of experience -> I have 1 year of experience
    text = re.sub(r'\b(?:one|two|three|four|five|six|seven|eight|nine|ten)\b', '', text, flags=re.IGNORECASE) 

    #removing stop words
    tokens = nltk.word_tokenize(text.lower()) #tokenize and convert to lower case
    tokens = [word for word in tokens if word not in stop_words] 

    return ' '.join(tokens)

# function to clean text from skills, degrees, designation
def clean_text(skills):
    skills = str(skills).replace('[','').replace(']','').replace("'",'')
    skills = skills.replace(' ','').lower()
    skills = skills.replace(',',' ')
    return skills

# using pyresparser library to extract features from resume
def extract_info(text):
    # creating a new Word document
    doc = docx.Document()
    doc.add_paragraph(text)
    doc.save("temp.docx")

    # using PyResparser to extract information from the resume
    extracted_info = pyresparser.ResumeParser("temp.docx").get_extracted_data()

    # extracting name, email and skills from the extracted information
    name = extracted_info['name']
    email = extracted_info['email']
    skills = extracted_info['skills']

    return name, email, skills

# Load the SVM model
filename = 'SVM_model.sav'
load_model = pickle.load(open(filename, 'rb'))

# Initialize the Flask application
app = Flask(__name__)

directory_path = r"uploads"

@app.route('/')
def index():
    return render_template('index.html')
@app.route('/predict', methods=['POST'])
def predict():
    files = request.files.getlist('file')

    job_desc = request.form['job_description']

    doc = docx.Document()
    doc.add_paragraph(job_desc)
    doc.save(f"job.docx")
    job_skills = pyresparser.ResumeParser(f"job.docx").get_extracted_data()['skills']
    cleaned_job_skills = clean_text(', '.join(job_skills))


    vectorizer = TfidfVectorizer()

    job_constraints_str = cleaned_job_skills


    job_constraints_vec = vectorizer.fit_transform([job_constraints_str.lower()])


    results = []
    for file in files:

        file.save(os.path.join(directory_path, file.filename))

        text = docx2txt.process(os.path.join(directory_path, file.filename))

        cleaned_text = clean_resume(text)


        name, email, skills = extract_info(text)

        cleaned_skills = clean_text(skills)
        prediction = load_model.predict([cleaned_text])[0]

        cleaned_resume_skills = clean_text(', '.join(skills))
        
        candidate_skills_vecs = vectorizer.transform([cleaned_resume_skills])


        cos_sim_scores = cosine_similarity(job_constraints_vec, candidate_skills_vecs).flatten()[0]
        

        results.append({'name': name, 'email': email, 'prediction': prediction, 'cos_sim_scores': cos_sim_scores,})
        
        results_sorted = sorted(results, key=lambda k: k['cos_sim_scores'], reverse=True)

    return render_template('index.html', results=results_sorted)





# Run the Flask application
if __name__ == '__main__':
    app.run(debug=True)
